# app.py - Flask Web Application for CBSE Q&A Generator
import os
import json
import random
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, session
from werkzeug.utils import secure_filename
import tempfile
from datetime import datetime

# Import your existing modules
from manifest_loader import ManifestLoader
from config import Config

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'  # Change this!

# Initialize config and manifest loader
config = Config()
manifest_loader = ManifestLoader()

# Global variables for loaded data
loaded_questions = {}
current_session = {}

@app.route('/')
def home():
    """Home page - shows subject selection"""
    return render_template('index.html', subjects=config.SUBJECTS)

@app.route('/api/subjects')
def get_subjects():
    """API endpoint to get available subjects"""
    return jsonify(config.SUBJECTS)

@app.route('/api/lessons/<subject>')
def get_lessons(subject):
    """Get lessons for a specific subject"""
    try:
        lessons = manifest_loader.get_lessons_for_subject(subject)
        return jsonify(lessons)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/question-types')
def get_question_types():
    """Get available question types"""
    return jsonify(config.QUESTION_TYPE_NAMES)

@app.route('/generator')
def generator():
    """Question generator page"""
    return render_template('generator.html', 
                         subjects=config.SUBJECTS,
                         question_types=config.QUESTION_TYPE_NAMES)

@app.route('/api/generate-questions', methods=['POST'])
def generate_questions():
    """Generate questions based on parameters"""
    try:
        data = request.json
        subject = data.get('subject')
        lessons = data.get('lessons', [])
        question_types = data.get('question_types', [])
        marks = data.get('marks', 20)
        
        # Generate questions using your existing logic
        questions = generate_question_set(subject, lessons, question_types, marks)
        
        # Store in session for download
        session['generated_questions'] = questions
        session['generation_params'] = data
        
        return jsonify({
            'success': True,
            'questions': questions,
            'total_marks': sum(q.get('Marks', 1) for q in questions)
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/practice')
def practice():
    """Practice mode page"""
    return render_template('practice.html', 
                         subjects=config.SUBJECTS,
                         question_types=config.QUESTION_TYPE_NAMES)

@app.route('/api/start-practice', methods=['POST'])
def start_practice():
    """Start a practice session"""
    try:
        data = request.json
        subject = data.get('subject')
        lesson = data.get('lesson')
        question_type = data.get('question_type')
        count = data.get('count', 10)
        
        # Load questions for practice
        questions = load_practice_questions(subject, lesson, question_type, count)
        
        # Create practice session
        session_id = f"practice_{datetime.now().timestamp()}"
        session[session_id] = {
            'questions': questions,
            'current_index': 0,
            'score': 0,
            'answers': []
        }
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'total_questions': len(questions),
            'first_question': questions[0] if questions else None
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/practice/<session_id>/next')
def get_next_question(session_id):
    """Get next question in practice session"""
    if session_id not in session:
        return jsonify({'error': 'Session not found'}), 404
    
    practice_session = session[session_id]
    current_index = practice_session['current_index']
    questions = practice_session['questions']
    
    if current_index >= len(questions):
        return jsonify({'finished': True, 'score': practice_session['score']})
    
    question = questions[current_index]
    return jsonify({
        'question': question,
        'index': current_index + 1,
        'total': len(questions)
    })

@app.route('/api/practice/<session_id>/answer', methods=['POST'])
def submit_answer(session_id):
    """Submit answer for current question"""
    if session_id not in session:
        return jsonify({'error': 'Session not found'}), 404
    
    practice_session = session[session_id]
    data = request.json
    user_answer = data.get('answer')
    
    current_index = practice_session['current_index']
    question = practice_session['questions'][current_index]
    correct_answer = question.get('Answer', '')
    
    # Simple scoring logic - you can make this more sophisticated
    is_correct = user_answer.lower().strip() == correct_answer.lower().strip()
    if is_correct:
        practice_session['score'] += 1
    
    practice_session['answers'].append({
        'question': question,
        'user_answer': user_answer,
        'correct_answer': correct_answer,
        'is_correct': is_correct
    })
    
    practice_session['current_index'] += 1
    session[session_id] = practice_session
    
    return jsonify({
        'correct': is_correct,
        'correct_answer': correct_answer,
        'explanation': question.get('Explanation', '')
    })

@app.route('/download/questions')
def download_questions():
    """Download generated questions as Word document"""
    if 'generated_questions' not in session:
        return "No questions to download", 400
    
    try:
        questions = session['generated_questions']
        params = session.get('generation_params', {})
        
        # Create temporary file
        temp_file = create_word_document(questions, params)
        
        return send_file(
            temp_file,
            as_attachment=True,
            download_name=f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return f"Error creating document: {str(e)}", 500

@app.route('/membership-check')
def membership_check():
    """Check membership status (integrate with your WordPress site)"""
    # This will integrate with your WordPress membership system
    # For now, return a placeholder
    return jsonify({
        'has_access': True,  # Replace with actual membership check
        'membership_type': 'student',  # or 'teacher'
        'expires': '2025-12-31'
    })

# Helper functions
def generate_question_set(subject, lessons, question_types, target_marks):
    """Generate a set of questions based on parameters"""
    all_questions = []
    
    for lesson in lessons:
        for q_type in question_types:
            try:
                questions = manifest_loader.load_questions_by_type_and_lesson(q_type, subject, lesson)
                all_questions.extend(questions)
            except Exception as e:
                print(f"Error loading {q_type} questions for {lesson}: {e}")
    
    # Simple selection logic - you can enhance this
    if not all_questions:
        return []
    
    # Randomly select questions up to target marks
    random.shuffle(all_questions)
    selected_questions = []
    current_marks = 0
    
    for question in all_questions:
        q_marks = question.get('Marks', 1)
        if current_marks + q_marks <= target_marks:
            selected_questions.append(question)
            current_marks += q_marks
        
        if current_marks >= target_marks:
            break
    
    return selected_questions

def load_practice_questions(subject, lesson, question_type, count):
    """Load questions for practice mode"""
    try:
        questions = manifest_loader.load_questions_by_type_and_lesson(question_type, subject, lesson)
        if len(questions) > count:
            questions = random.sample(questions, count)
        return questions
    except Exception as e:
        print(f"Error loading practice questions: {e}")
        return []

def create_word_document(questions, params):
    """Create Word document with questions"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add header
        title = doc.add_heading('CBSE Social Science Question Paper', 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        doc.add_paragraph(f"Subject: {params.get('subject', 'Social Science')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph(f"Total Marks: {sum(q.get('Marks', 1) for q in questions)}")
        
        doc.add_paragraph("").add_run().add_break()
        
        # Add questions
        for i, question in enumerate(questions, 1):
            q_para = doc.add_paragraph()
            q_para.add_run(f"Q{i}. ").bold = True
            q_para.add_run(question.get('Text', ''))
            
            # Add marks
            marks = question.get('Marks', 1)
            q_para.add_run(f" [{marks} mark{'s' if marks > 1 else ''}]").italic = True
            
            # Add options for MCQ
            if question.get('Type') == 'mcq' and question.get('Options'):
                for option in question['Options']:
                    doc.add_paragraph(f"  {option}", style='List Bullet')
            
            doc.add_paragraph("")  # Space between questions
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        raise

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)